import wx
import wx.richtext as richtext
from formeqns import EqnForm
import os
import GTC as gtc
import docx
from docx.shared import Pt
import time
from PIL import Image
import io


class REPORT(EqnForm):

    def __init__(self, parent):
        """
        All the writing processes for the final reporting
        """
        EqnForm.__init__(self, parent)

    def OnSave(self, event):
        """
        Saves a Word version of the report with a suggested folder and file name.
        The default is to save it in the same folder as the Excel spread sheet
        while appending .docx to the full Excel file name.
        """
        wildcard = "Project report (*.docx)|*.docx|" \
                   "All files (*.*)|*.*"
        dirname = self.projwd[:-9]  # the one above mie_temp (remove 9 characters)
        dlg = wx.FileDialog(self, "Choose folder to save report", defaultDir=dirname,
                            defaultFile=self.proj_file + '.docx',
                            wildcard=wildcard, style=wx.FD_SAVE)

        if dlg.ShowModal() == wx.ID_OK:
            filename = dlg.GetFilename()
            dirname = dlg.GetDirectory()
            docx_file = os.path.join(dirname, filename)
            # create and save word report
            self.wordreport(docx_file)
        dlg.Destroy()

    def OnPrint(self, event):
        self.PrintNow()

    def PrintNow(self):
        """
        Prints the rich text report without preview but with a printer dialog.
        Printing to a system pdf printer can be a convenient way of saving a
        record of the report.
        """
        rtp = richtext.RichTextPrinting("Print")
        # rtp.PageSetup()  # should be able to change print margins without dialogue?
        localtime = time.asctime(time.localtime(time.time()))
        footer = self.m_statusBar1.GetStatusText(1) + ' printed on ' + localtime
        rtp.SetFooterText(footer)  # reads file name
        rtp.PrintBuffer(self.report_richText.GetBuffer())

    def both(self, row, rich, value):
        """
        This simultaneously sends *value* to *row* to write to the file and
        writes the same *value* to the *rich* text window.  There should be a
        more elegant way to do this.
        """
        rich.AppendText(value)
        rich.AppendText(', ')
        row.append(value)

    def reporter(self, error, temperature):
        """
        Creates strings, including numerical values calculated from the GTC
        ureal final_error, for feeding into reports. It creates a list of
        png files for inclusion in reports.
        :param error: is a list of the errors as calculated in grand_finale
        :param temperature: is the ureal temperature value for the site
        :return:
        """

        self.report_txt = []  # vital to clear these lists before writing a new report
        self.report_images = []
        self.report_images_wx = []
        final_error = []
        uncertainty = []
        maximum = []
        minimum = []
        for e in error:  # create lists by load profile of the error
            final_error.append(e.x)
            uncert = e.u * gtc.reporting.k_factor(e.df)
            uncertainty.append(uncert)
            maximum.append(e.x + uncert)
            minimum.append(e.x - uncert)

        # summarising uncertainty
        meterCal = ['a0 meter', 'a1 meter', 'a2 meter', 'a3 meter', 'b0 meter', 'Meter annual drift']
        meterInf = ['Meter field dependence', 'Meter temperature coefficient', 'Meter frequency dependence',
                    'Meter voltage dependence', 'Meter harmonic dependence']
        ctCal = ['a0 CT_e', 'a0 CT_p', 'a1 CT_e', 'a1 CT_p', 'a2 CT_e', 'a2 CT_p', 'b0 CT_e', 'b0 CT_p',
                 'CT calibration temperature', 'CT calibration burden']
        ctInf = ['CT temperature coefficient of error', 'CT temperature coefficient of phase',
                 'CT burden coefficient of error', 'CT burden coefficient of phase']
        vtCal = ['a0 VT_e', 'a0 VT_p', 'a1 VT_e', 'a1 VT_p', 'a2 VT_e', 'a2 VT_p', 'b0 VT_e', 'b0 VT_p',
                 'VT calibration temperature', 'VT calibration burden']
        vtInf = ['VT temperature coefficient of error', 'VT temperature coefficient of phase',
                 'VT burden coefficient of error', 'VT burden coefficient of phase']
        siteInf = ['Site EM field', 'Site temperature', 'Site voltage', 'Site frequency', 'Site harmonics',
                   'Site VT burden', 'Site CT burden']
        complist = [meterCal, meterInf, ctCal, ctInf, vtCal, vtInf, siteInf]

        normalised_share = []
        for i in range(self.n_profiles):
            xx = self.model.budget_by_label(error[i], complist)
            excluding_site = xx[0] + xx[1] + xx[2] + xx[3] + xx[4] + xx[5]
            share = [xx[0] / excluding_site, xx[1] / excluding_site, xx[2] / excluding_site,
                                xx[3] / excluding_site, xx[4] / excluding_site, xx[5] / excluding_site]
            normalised_share.append(share)
        # prepare text/strings needed for any output version
        self.report_txt.append('Metering Installation Error Report')
        self.report_txt.append('Accuracy')
        plural = ''  # no s if singular
        if self.n_profiles > 1:
            plural = 's'  # add s for plural
        nxtln = 'The accuracy of the metering installation (combined meter and transformers) has been calculated for ' \
                   + str(self.n_profiles) + ' annual load profile' + plural + '.'
        self.report_txt.append(nxtln)  # 2
        self.report_txt.append('Load Profile\t\tError\t\t\tUncertainty\t\tError-Uncertainty\tError+Uncertainty')
        self.word_tabs = []  # differently tabbed version for word
        for load_no in range(self.n_profiles):
            _err = f"{final_error[load_no]:.2f}"  # formatting float into a 2-decimal-place string
            _unc = f"{uncertainty[load_no]:.2f}"
            _min = f"{minimum[load_no]:.2f}"
            _max = f"{maximum[load_no]:.2f}"
            nxtln = '      ' + str(load_no + 1) + '\t\t\t' + _err +' %\t\t' + _unc + ' %\t\t' + _min + ' %\t\t'\
                    + _max + ' %'
            word_ln = '      ' + str(load_no + 1) + '\t\t\t'
            if _err[0] == '-':
                word_ln = word_ln + _err + ' %\t'
            else:
                word_ln = word_ln + _err +' %\t\t'
            word_ln = word_ln + _unc + ' %\t\t'
            if _min[0] == '-':
                word_ln = word_ln + _min + ' %\t' + _max + ' %'
            else:
                word_ln = word_ln + _min + ' %\t\t' + _max + ' %'
            self.report_txt.append(nxtln)
            self.word_tabs.append(word_ln)
            # self.report_txt.append(word_ln)
        self.report_txt.append('The uncertainty is the expanded uncertainty calculated at a 95 % level of confidence. ')
        self.report_txt.append('For each load profile the same variation in temperature and network conditions was '
                               'assumed. Sensitivity coefficients published for an IEC class ' + self.iec +
                               ' meter were incorporated in the uncertainty calculation.')
        # temp_max = f"{temperature.x + 2 * temperature.u:.1f}" + ' ' + u'\N{DEGREE SIGN}' + 'C'
        # temp_min = f"{temperature.x - 2 * temperature.u:.1f}" + ' ' + u'\N{DEGREE SIGN}' + 'C'
        mean_temp = f"{temperature.x:.1f}" + ' ' + u'\N{DEGREE SIGN}' + 'C'
        u_temp = f"{temperature.u * 2.0:.1f}" + ' ' + u'\N{DEGREE SIGN}' + 'C'
        # self.report_txt.append('The uncertainty includes a contribution due to the metering installation temperature '
        #                        'varying over the range of ' + temp_min + ' to ' + temp_max + ' during the year.')
        self.report_txt.append('The uncertainty includes a contribution due to the temperature of the metering '
                               'installation differing from the temperature at which the components were calibrated. An'
                               ' average value of ' + mean_temp + ' with an expanded uncertainty of ' + u_temp + ' is '
                               'assumed for a twelve-month period.')
        self.report_txt.append('Annual Load Profile' + plural)
        self.report_txt.append('Annual load profile' + plural + ' used for this error calculation. The graph is '
                                                                'normalised to give the relative amount of energy '
                                                                'at each combination of current and phase angle.')

        self.report_txt.append('Error Plot' + plural)
        self.report_txt.append('Installation error over the phase and current values for the load profile' + plural)
        self.report_txt.append('Uncertainty by Component')
        self.report_txt.append('Uncertainty contribution given by component as a percentage of total variance. This is '
                               'calculated for each load profile. This may help with any decisions about which '
                               'components might be considered for an upgrade')
        self.report_txt.append('Load Profile\t\tMeter\t\tCT\tVT')
        for load_no in range(self.n_profiles):
            _meter = f"{(normalised_share[load_no][0] + normalised_share[load_no][1]) * 100:.0f}"  # formatting float into a 0-decimal-place string
            _ct = f"{(normalised_share[load_no][2] + normalised_share[load_no][3]) * 100:.0f}"
            _vt = f"{(normalised_share[load_no][4] + normalised_share[load_no][5]) * 100:.0f}"
            nxtln = '      ' + str(load_no + 1) + '\t\t\t' + _meter +' %\t\t' + _ct + ' %\t' + _vt + ' %'
            self.report_txt.append(nxtln)
        self.report_txt.append('Component Calibration')
        self.report_txt.append('For each component the calibration points are plotted on a graph with uncertainty '
                               'bars. Fitted curves are shown with an upper and lower curve indicating the uncertainty '
                               ' of the fit. The coefficients of these curves are used to calculate the metering '
                               'installation error.')

        self.report_txt.append('Current Transformer')
        self.report_txt.append('Voltage Transformer')
        self.report_txt.append('Meter')
        self.report_txt.append('Installation Error Contour')
        self.report_txt.append('Additional Information')

        # gather graphs for reporting
        image = self.buffer_image_wx(self.report_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.load_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.CT_graph_1.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.VT_graph_1.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.meter_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.report_contour.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])


    def buffer_image_wx(self, plot_canvas):
        """

        Takes the plot on a canvas and prints it to a buffer which is then scaled
        and created as a wx richtext compatible image. The size (360, 270) could be
        an input parameter.
        :param plot_canvas: is a canvas in wx holding a graph
        :return:
        """
        buffer = io.BytesIO()
        plot_canvas.print_figure(buffer, format='png', dpi=300)
        buffer.seek(0)
        img1 = Image.open(buffer)
        img = img1.resize((360, 270), Image.LANCZOS)
        image = wx.Image(img.size[0], img.size[1])
        image.SetData(img.convert("RGB").tobytes())
        return image, buffer

    def heading(self, txt, font):
        """
        the txt string is written to self.report_richText in bold with the chosen font size
        :param txt: string to be written
        :param font: font point size as integer
        :return:
        """
        report = self.report_richText
        report.BeginFontSize(font)
        report.BeginBold()
        report.WriteText(txt)
        report.Newline()
        report.EndFontSize()
        report.EndBold()

    def wxreport(self):
        """
        Takes the *t_strings* and *image_files* lists prepared in 'reporter'
        and writes the report to the wxRichText panel in the GUI.
        """
        t_strings = self.report_txt
        # write to the wxRichText panel
        report = self.report_richText
        self.heading(t_strings[0],14)
        self.heading(t_strings[1], 12)
        report.BeginFontSize(10)
        report.WriteText(t_strings[2])
        report.Newline()
        report.BeginBold()
        report.WriteText(t_strings[3])
        report.Newline()
        report.EndBold()
        for i in range(self.n_profiles):
            report.WriteText(t_strings[i + 4])
            report.Newline()
        index = i + 4  # assumes 4 lines before the variable table
        report.WriteText(t_strings[index +1])
        report.WriteText(t_strings[index + 2])
        report.Newline()
        report.WriteText(t_strings[index + 3])
        report.EndFontSize()
        report.Newline()
        self.heading(t_strings[index + 4], 12)
        report.BeginFontSize(10)  # Beginning of text for load profiles
        report.WriteText(t_strings[index + 5])
        report.Newline()
        report.WriteImage(self.report_images_wx[1])  # the load profile
        report.Newline()
        report.Newline()
        report.EndFontSize()
        self.heading(t_strings[index + 6], 12)
        report.BeginFontSize(10)
        report.WriteText(t_strings[index + 7])
        report.Newline()
        report.WriteImage(self.report_images_wx[0])  # the error plots
        report.Newline()
        report.Newline()
        report.EndFontSize()
        self.heading(t_strings[index + 8], 12)
        report.BeginFontSize(10)
        report.WriteText(t_strings[index + 9])
        report.Newline()
        report.BeginBold()
        report.WriteText(t_strings[index + 10])
        report.EndBold()
        report.Newline()
        for i in range(self.n_profiles):
            report.WriteText(t_strings[index + 11 + i])
            report.Newline()
        index = index + 11 + i  # account for variable number of load profile results
        report.Newline()
        report.EndFontSize()
        self.heading(t_strings[index + 1], 12)
        report.BeginFontSize(10)
        report.WriteText(t_strings[index + 2])
        report.Newline()
        report.BeginBold()
        report.WriteText(t_strings[index + 3])
        report.Newline()
        report.WriteImage(self.report_images_wx[2])  # CT graphs
        report.Newline()
        report.WriteText(t_strings[index + 4])
        report.Newline()
        report.WriteImage(self.report_images_wx[3])
        report.Newline()
        report.WriteText(t_strings[index + 5])  # meter
        report.Newline()
        report.WriteImage(self.report_images_wx[4])
        report.Newline()
        report.WriteText(t_strings[index + 6])  # contour
        report.Newline()
        report.WriteImage(self.report_images_wx[5])
        report.EndFontSize()

    def wordreport(self, docx_file):
        """
        Takes the *t_strings* and *image_files* lists prepared in 'reporter' and writes
        the report to a word document. The document will display the text in the
        default styles in the Word installation used to view the file.  The image
        files retain their native resolution. By default, the name and folder of
        the Excel input file is used simply with .docx added at the end. A folder
        'templates' is assumed to hold default.docx on which to build the report.
        """
        t_strings = self.report_txt
        assert len(t_strings) > 0, "Nothing to report.  Run calculation before attempting to save"
        template = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'templates', 'default.docx')
        document = docx.Document(template)
        section = document.sections[0]
        footer = section.footer
        localtime = time.asctime(time.localtime(time.time()))
        identifier = '(Data source: ' + self.m_statusBar1.GetStatusText(1) + '. Error calculated: ' + localtime + '.)'
        paragraph = footer.paragraphs[0]
        paragraph.text = identifier
        document.add_heading(t_strings[0], level=0)
        document.add_heading(t_strings[1], level=1)
        document.add_paragraph(t_strings[2])
        # document.add_paragraph(t_strings[3])  # force a differently tabbed header
        nxtln = 'Load Profile\t\tError\t\tUncertainty\tError-Uncert\tError+Uncert'
        document.add_paragraph(nxtln)
        for i in range(self.n_profiles):
            # document.add_paragraph(t_strings[i + 4])
            document.add_paragraph(self.word_tabs[i])  # use differently tabbed version
        index = i + 4  # assumes 4 lines before the variable table
        document.add_paragraph(t_strings[index + 1] + t_strings[index + 2])
        document.add_paragraph (t_strings[index + 3])
        document.add_heading(t_strings[index + 4], level=1)
        document.add_paragraph(t_strings[index + 5])
        image_files = self.report_images
        document.add_picture(image_files[1], width=docx.shared.Mm(100))
        document.add_heading(t_strings[index + 6], level=1)
        document.add_paragraph(t_strings[index + 7])
        document.add_picture(image_files[0], width=docx.shared.Mm(100))
        document.add_heading(t_strings[index + 8], level=1)
        document.add_paragraph(t_strings[index + 9])
        # document.add_paragraph(t_strings[index + 10])  # force a differently tabbed header
        nxtln = 'Load Profile\t\tMeter\t\tCT\tVT'
        document.add_paragraph(nxtln)
        for i in range(self.n_profiles):
            document.add_paragraph(t_strings[index + 11 + i])
        index = i + index + 11
        document.add_page_break()
        document.add_heading(t_strings[index + 1])
        document.add_paragraph(t_strings[index + 2])
        document.add_heading(t_strings[index + 3], level=1)
        document.add_picture(image_files[2], width=docx.shared.Mm(100))
        document.add_heading(t_strings[index + 4], level=1)
        document.add_picture(image_files[3], width=docx.shared.Mm(100))
        document.add_heading(t_strings[index + 5], level=1)
        document.add_picture(image_files[4], width=docx.shared.Mm(100))
        document.add_heading(t_strings[index + 6], level=1)
        document.add_picture(image_files[5], width=docx.shared.Mm(100))
        document.add_page_break()
        document.add_heading('Additional Information', level=1)
        document.add_heading('Text Output From the Calculation Process', level=2)
        no_lines = self.m_textCtrl3.GetNumberOfLines()
        for x in range(no_lines):
            output_notes = self.m_textCtrl3.GetLineText(x)
            paragraph = document.add_paragraph(output_notes)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(2)
            paragraph_format.space_after = Pt(2)
            # document.add_paragraph(output_notes)
        document.save(docx_file)


if __name__ == '__main__':
    app = wx.App()
    frame = REPORT(None)
    frame.Show()
    app.MainLoop()