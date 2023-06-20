import wx
import wx.richtext as richtext
from formeqns import EqnForm
import os
import csv
import components as comp
import functions
import numpy as np
import warnings
from scipy.stats import t
import GTC as gtc
import extras
import docx
import time
import sys
from PIL import Image
import io


class MIECALC(EqnForm):

    def __init__(self, parent):
        """
        Adds the overall calculation process.
        """
        EqnForm.__init__(self, parent)
        log = self.m_textCtrl3  # where stdout will be redirected
        redir = extras.RedirectText(log)
        sys.stdout = redir
        self.m_button26.Enable(False)  # 'Process project file' button initially greyed out


    #######################Creating Summary Files###########################
    def OnMeterSummary(self, event):
        self.PushMeterSummary()

    def PushMeterSummary(self):
        """
        Button to create csv summary file for meter.
        """
        self.create_summary(self.eqn_choice_Meter, self.meter_table, self.Meter_fit_grid, self.meter_richText, 'meter',
                            self.e_data('_meter.csv'))

    def both(self, row, rich, value):
        """
        This simultaneously sends *value* to *row* to write to the file and
        writes the same *value* to the *rich* text window.  There should be a
        more elegant way to do this.
        """
        rich.AppendText(value)
        rich.AppendText(', ')
        row.append(value)

    def create_site_inf_summary(self, table, text, sfile):
        """
        This is a special case of *create_summary* as the site only has influence
        variables.  It could be made a conditional subset of create_summary.
        Creates a csv *sfile* from the data in *table* and also displays it in the
        relevant *text* window.
        """
        writer = csv.writer(open(sfile, 'w', newline=''), 'excel')
        row = []
        text.Clear()
        text.AppendText(sfile)
        row.append(sfile)
        writer.writerow(row)
        text.Newline()
        row = []
        table_rows = table.GetNumberRows()
        for i in range(table_rows):
            self.both(row, text, str(table.GetCellValue(i, 1)))
            self.both(row, text, (str(table.GetCellValue(i, 3))))
            self.both(row, text, str(table.GetCellValue(i, 5)))
            self.both(row, text, table.GetCellValue(i, 0))
            text.Newline()
            writer.writerow(row)
            row = []

    def create_summary(self, choice, table, grid, text, component, cfile):
        """
        Primarily produces a csv file of coefficients and influences for use
        in COMPONENT classes.  Fitting coefficients are from *grid* and
        influence data is from *table*.The content of the csv file, *cfile* is
        echoed in the rich text box, *text*.  Note that this is where the degrees
        of freedom for *b0*, the bias coefficient, are set to 50.
        """
        writer = csv.writer(open(cfile, 'w', newline=''), 'excel')
        row = []
        text.Clear()
        text.AppendText(cfile)
        row.append(cfile)
        writer.writerow(row)
        text.Newline()
        row = []
        text.AppendText(str(choice.GetSelection()))
        row.append(choice.GetSelection())
        writer.writerow(row)
        text.Newline()
        row = []
        gridrows = grid.GetNumberRows()
        no_coeffs = int((gridrows - 5) / 2.0)
        for i in range(no_coeffs):
            self.both(row, text, grid.GetCellValue(i + 1, 2))
            self.both(row, text, grid.GetCellValue(i + 1, 3))
            self.both(row, text, grid.GetCellValue(no_coeffs + 2, 1))
            self.both(row, text, grid.GetCellValue(i + 1, 0) + ' ' + component)
            text.Newline()
            writer.writerow(row)
            row = []
        # get the covariance matrix
        a = no_coeffs + 5  # a and b are grid offsets
        b = 1
        for i in range(no_coeffs):
            for j in range(no_coeffs):
                self.both(row, text, grid.GetCellValue(a + i, b + j))

        text.Newline()
        writer.writerow(row)
        row = []
        self.both(row, text, str(grid.GetCellValue(no_coeffs + 1, 2)))
        self.both(row, text, str(grid.GetCellValue(no_coeffs + 1, 3)))
        self.both(row, text, '50')  # force d.o.f. for b0
        self.both(row, text, grid.GetCellValue(no_coeffs + 1, 0) + ' ' + component)
        text.Newline()
        writer.writerow(row)
        row = []
        table_rows = table.GetNumberRows()
        for i in range(table_rows):
            self.both(row, text, str(table.GetCellValue(i, 1)))
            self.both(row, text, (str(table.GetCellValue(i, 3))))
            self.both(row, text, str(table.GetCellValue(i, 5)))
            self.both(row, text, table.GetCellValue(i, 0))
            text.Newline()
            writer.writerow(row)
            row = []

    def JoinErrorPhase(self, error, phase, joined):
        """
        Concatenates the *error* and *phase* files produced by 'create_summary'
        into the single *joined* file that does not include the header row of the
        phase file.  This is overhead management to produce the required csv file.
        """
        reader1 = csv.reader(open(error, 'r'))
        reader2 = csv.reader(open(phase, 'r'))
        writer = csv.writer(open(joined, 'w', newline=''))
        rownum = 0
        for row in reader1:
            if rownum == 0:
                writer.writerow(str(self.e_data('_CT.csv')))  # why are characters separated??
            else:
                writer.writerow(row)
            rownum += 1
        end_row = rownum
        for row in reader2:
            if rownum == end_row:
                pass
            else:
                writer.writerow(row)
            rownum += 1

    def OnCTSummary(self, event):
        self.PushCTSummary()

    def PushCTSummary(self):
        """
        Button to automatically put all CT data into csv files.
        """
        self.create_summary(self.eqn_choice_CTratio, self.CT_table, self.CTratio_fit_grid, self.CT_richText, 'CT_e',
                            self.e_data('_CTe.csv'))
        self.create_summary(self.eqn_choice_CTphase, self.CT_table1, self.CTphase_fit_grid, self.CT_richText1, 'CT_p',
                            self.e_data('_CTp.csv'))
        self.JoinErrorPhase(self.e_data('_CTe.csv'), self.e_data('_CTp.csv'), self.e_data('_CT.csv'))

    def OnVTSummary(self, event):
        self.PushVTSummary()

    def PushVTSummary(self):
        """
        Button push to automatically put all VT data into csv files.
        """
        self.create_summary(self.eqn_choice_VTratio, self.VT_table, self.VTratio_fit_grid, self.VT_richText, 'VT_e',
                            self.e_data('_VTe.csv'))
        self.create_summary(self.eqn_choice_VTphase, self.VT_table1, self.VTphase_fit_grid, self.VT_richText1, 'VT_p',
                            self.e_data('_VTp.csv'))
        self.JoinErrorPhase(self.e_data('_VTe.csv'), self.e_data('_VTp.csv'), self.e_data('_VT.csv'))

    def OnSiteSummary(self, event):
        self.PushSiteSummary()

    def PushSiteSummary(self):
        """
        Button push to automatically put all CT data into csv files.
        """
        self.create_site_inf_summary(self.site_table, self.site_richText, self.e_data('_site.csv'))

    ####################End of Creating Summary Files########################

    ############Manage Calculation & Reports############
    def OnGUMcalc(self, event):
        self.PushGUMcalc()

    def PushGUMcalc(self):
        """
        Executes grand_finale assuming that all data has been loaded either from
        file or manually. This is run when the 'Process manually created files'
        button is pushed.
        """
        self.projwd = self.m_textCtrl999.GetValue()  # the project working directory can be altered
        self.m_staticText22.SetLabel('Project directory:  ' + self.projwd)  # display on 'Input file list and messages'
        self.m_statusBar1.SetStatusText('Calculation started', 2)
        self.m_statusBar1.SetStatusText('Meter calculation', 0)
        self.PushLoadMeterData()
        self.PushPlotMeter()
        print('Meter data')
        self.PushFit_Meter()
        self.PushLoadMeterInf()
        self.PushMeterSummary()
        self.m_statusBar1.SetStatusText('CT calculation', 0)
        self.PushLoadCTData()
        self.PushPlotCTphase()
        self.PushPlotCTratio()
        print('CT ratio data')
        self.PushFit_CTratio()
        print('CT phase data')
        self.PushFit_CTphase()
        self.PushLoadCTInf()
        self.PushCTSummary()
        self.m_statusBar1.SetStatusText('VT calculation', 0)
        self.PushLoadVTData()
        self.PushPlotVTphase()
        self.PushPlotVTratio()
        print('VT ratio data')
        self.PushFit_VTratio()
        print('VT phase data')
        self.PushFit_VTphase()
        print('')
        self.PushLoadVTInf()
        self.PushVTSummary()
        self.m_statusBar1.SetStatusText('Load processing', 0)
        self.PushPlotLoadProfile()
        self.PushLoadSiteInf()
        self.PushSiteSummary()
        self.m_statusBar1.SetStatusText('Overall calculation', 0)
        self.grand_finale()

    def grand_finale(self):
        """
        The error is calculated at the (current, phase) points provided in the
        load profile as given in load.csv.  A plot of the overall error against
        current and phase angle is created and a screen report that has the key
        results and graphs is generated.
        """
        # assemble components and calculate the error
        data1 = self.load_values.GetValue()
        data2 = self.load_data.GetValue()
        profile = comp.LOAD('default', data1, data2)  # data files as from load table
        model = functions.MODEL('model')  # pick up the standard model functions to pass into component objects
        meter = comp.METER('default', '1 element', model, self.e_data('_meter.csv'))
        ct = comp.TRAN('name', 'single', model, self.e_data('_CT.csv'))
        vt = comp.TRAN('name', 'single', model, self.e_data('_VT.csv'))
        site = comp.INSTALLATION('name', meter, ct, vt, profile, self.e_data('_site.csv'))
        error = site.site_error_terms()
        self.m_statusBar1.SetStatusText('Calculation finished', 2)

        # Plot the error
        r = error[2]  # installation's X as determined by the profile
        no_of_points = len(r)
        X = np.zeros(no_of_points)
        Y = np.zeros(no_of_points)
        Z = np.zeros(no_of_points)
        Z1 = np.zeros(no_of_points)
        Z2 = np.zeros(no_of_points)
        for i in range(no_of_points):
            X[i] = r[i][0]
            Y[i] = r[i][1]
            Z[i] = error[0][i].x
            uZ = error[0][i].u
            dfZ = error[0][i].df
            kZ = t.ppf(0.975, dfZ)  # note scipy t
            Z1[i] = Z[i] + uZ * kZ
            Z2[i] = Z[i] - uZ * kZ
        edge_no = np.sqrt(no_of_points)  # assuming an nxn grid
        dummy = np.zeros((int(edge_no), int(edge_no)), dtype=float)  # assumes nxn
        ZZ = np.reshape(Z, np.shape(dummy))
        XX = np.reshape(X, np.shape(dummy))
        YY = np.reshape(Y, np.shape(dummy))
        ZZ1 = np.reshape(Z1, np.shape(dummy))
        ZZ2 = np.reshape(Z2, np.shape(dummy))
        with warnings.catch_warnings():  # get 'converting masked element to nan'
            warnings.simplefilter("ignore")
            np.seterr(invalid='ignore')
            self.report_graph.ax.plot_surface(XX, YY, ZZ, rstride=1, cstride=1, cmap='jet')
            self.report_graph.ax.plot_wireframe(XX, YY, ZZ1)
            self.report_graph.ax.plot_wireframe(XX, YY, ZZ2)
            np.seterr(invalid='print')
        self.report_graph.ax.autoscale(enable=True, axis='both', tight=True)
        self.report_graph.canvas.draw()
        # create report
        self.m_statusBar1.SetStatusText('Generating report', 0)
        self.reporter(error[1])
        self.wxreport()

    def OnSave(self, event):
        """
        Saves a Word version of the report with a suggested folder and file name.
        The default is to save it in the same folder as the Excel spread sheet
        while appending .docx to the full Excel file name.
        """
        wildcard = "Project report (*.docx)|*.docx|" \
                   "All files (*.*)|*.*"
        dirname = self.projwd[:-9]  # the one above xls_temp (remove 9 characters)
        dlg = wx.FileDialog(self, "Choose folder to save report", defaultDir=dirname,
                            defaultFile=self.proj_file + '.docx',
                            wildcard=wildcard, style=wx.FD_SAVE)

        if dlg.ShowModal() == wx.ID_OK:
            filename = dlg.GetFilename()
            dirname = dlg.GetDirectory()
            docx_file = os.path.join(dirname, filename)
            # create and save word report
            self.wordreport(docx_file)
        dlg.Destroy()

    def OnPrint(self, event):
        self.PrintNow()

    def PrintNow(self):
        """
        Prints the rich text report without preview but with a printer dialog.
        Printing to a system pdf printer can be a convenient way of saving a
        record of the report.
        """
        rtp = richtext.RichTextPrinting("Print")
        localtime = time.asctime(time.localtime(time.time()))
        footer = self.m_statusBar1.GetStatusText(1) + ' printed on ' + localtime
        rtp.SetFooterText(footer)  # reads file name
        rtp.PrintBuffer(self.report_richText.GetBuffer())

    def reporter(self, final_error):
        """
        Creates strings, including numerical values calculated from the GTC
        ureal *final_error*, for feeding into reports. It also returns a list of
        png files for inclusion in reports.
        """
        self.report_txt = []  # vital to clear these lists before writing a new report
        self.report_images = []
        self.report_images_wx = []
        # overall_error = final_error
        uncertainty = final_error.u * gtc.reporting.k_factor(final_error.df)
        maximum = final_error.x + uncertainty
        minimum = final_error.x - uncertainty

        # summarising uncertainty
        meterCal = ['a0 meter', 'a1 meter', 'a2 meter', 'a3 meter', 'b0 meter', 'Meter annual drift']
        meterInf = ['Meter field dependence', 'Meter temperature coefficient', 'Meter frequency dependence',
                    'Meter voltage dependence', 'Meter harmonic dependence']
        ctCal = ['a0 CT_e', 'a0 CT_p', 'a1 CT_e', 'a1 CT_p', 'a2 CT_e', 'a2 CT_p', 'b0 CT_e', 'b0 CT_p',
                 'CT calibration temperature', 'CT calibration burden']
        ctInf = ['CT temperature coefficient of error', 'CT temperature coefficient of phase',
                 'CT burden coefficient of error', 'CT burden coefficient of phase']
        vtCal = ['a0 VT_e', 'a0 VT_p', 'a1 VT_e', 'a1 VT_p', 'a2 VT_e', 'a2 VT_p', 'b0 VT_e', 'b0 VT_p',
                 'VT calibration temperature', 'VT calibration burden']
        vtInf = ['VT temperature coefficient of error', 'VT temperature coefficient of phase',
                 'VT burden coefficient of error', 'VT burden coefficient of phase']
        siteInf = ['Site EM field', 'Site temperature', 'Site voltage', 'Site frequency', 'Site harmonics',
                   'Site VT burden', 'Site CT burden']
        complist = [meterCal, meterInf, ctCal, ctInf, vtCal, vtInf, siteInf]
        xx = self.model.budget_by_label(final_error, complist)
        excluding_site = xx[0] + xx[1] + xx[2] + xx[3] + xx[4] + xx[5]
        normalised_share = [xx[0] / excluding_site, xx[1] / excluding_site, xx[2] / excluding_site,
                            xx[3] / excluding_site, xx[4] / excluding_site, xx[5] / excluding_site]

        # prepare text/strings needed for any output version
        self.report_txt.append('Metering Installation Error Report')  # 0
        self.report_txt.append('The error calculated for the given load profile is ')  # 1
        self.report_txt.append("%+ 0.2f %s" % (final_error.x, '%.'))  # 2
        self.report_txt.append('The expanded uncertainty at an estimated 95% level of confidence is')  # 3
        self.report_txt.append("%s %0.2f %s" % ('', uncertainty, '%.'))  # 4
        self.report_txt.append('The uncertainty interval extends from ')  # 5
        self.report_txt.append("%+ 0.2f %s to %+0.2f %s" % (minimum, '%', maximum, '%.'))  # 6
        self.report_txt.append('Uncertainty contribution by component as a percentage of total variance,')  # 7
        self.report_txt.append('Meter:')  # 8
        self.report_txt.append("%s %0.2G %s" % (' ', (normalised_share[0] + normalised_share[1]) * 100, ' %'))  # 9
        self.report_txt.append('CT:')  # 10
        self.report_txt.append("%s %0.2G %s" % (' ', (normalised_share[2] + normalised_share[3]) * 100, ' %'))  # 11
        self.report_txt.append('VT:')  # 12
        self.report_txt.append("%s %0.2G %s" % (' ', (normalised_share[4] + normalised_share[5]) * 100, ' %'))  # 13
        self.report_txt.append('A full list of components contributing to the total uncertainty can be found in the')  # 14
        self.report_txt.append('report workbook tab, "Input file list and messages".')  # 15
        self.report_txt.append('Installation error over the phase and current ranges that match the load profile.')  # 16
        self.report_txt.append('Annual load profile used for this error calculation.')  # 17
        self.report_txt.append(
            'Component calibration data with calculated fits over the calibration ranges shown overleaf.')  # 18
        self.report_txt.append('Current Transformer')  # 19
        self.report_txt.append('Voltage Transformer')  # 20
        self.report_txt.append('Meter')  # 21

        # gather graphs for reporting
        image = self.buffer_image_wx(self.report_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.load_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.CT_graph_1.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.VT_graph_1.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])
        image = self.buffer_image_wx(self.meter_graph.canvas)
        self.report_images_wx.append(image[0])
        self.report_images.append(image[1])

    def buffer_image_wx(self, plot_canvas):
        """

        Takes the plot on a canvas and prints it to a buffer which is then scaled
        and created as a wx richtext compatible image. The size (360, 270) could be
        an input parameter.
        :param plot_canvas: is a canvas in wx holding a graph
        :return:
        """
        buffer = io.BytesIO()
        plot_canvas.print_figure(buffer, format='png', dpi=300)
        buffer.seek(0)
        img1 = Image.open(buffer)
        img = img1.resize((360, 270), Image.LANCZOS)
        image = wx.Image(img.size[0], img.size[1])
        image.SetData(img.convert("RGB").tobytes())
        return image, buffer

    def wxreport(self):
        """
        Takes the *t_strings* and *image_files* lists prepared in 'reporter'
        and writes the report to the wxRichText panel in the GUI.
        """
        t_strings = self.report_txt
        # write to the wxRichText panel
        report = self.report_richText
        report.BeginFontSize(14)
        report.BeginBold()
        report.WriteText(t_strings[0])
        report.Newline()
        report.EndBold()
        report.EndFontSize()
        report.BeginFontSize(10)
        report.WriteText(t_strings[1])
        report.BeginBold()
        report.WriteText(t_strings[2])
        report.EndBold()
        report.Newline()
        report.WriteText(t_strings[3])
        report.BeginBold()
        report.WriteText(t_strings[4])
        report.EndBold()
        report.Newline()
        report.WriteText(t_strings[5])
        report.WriteText(t_strings[6])
        report.Newline()
        report.Newline()
        report.WriteText(t_strings[7])
        report.Newline()
        report.WriteText(t_strings[8])
        report.WriteText('\t')
        report.WriteText(t_strings[9])
        report.Newline()
        report.WriteText(t_strings[10])
        report.WriteText('\t\t')
        report.WriteText(t_strings[11])
        report.Newline()
        report.WriteText(t_strings[12])
        report.WriteText('\t\t')
        report.WriteText(t_strings[13])
        report.Newline()
        report.WriteText(t_strings[14])
        report.Newline()
        report.WriteText(t_strings[15])
        report.Newline()
        report.Newline()

        # add graph images to report, first rescaling them to fit a simple printout.
        # image resolution is degraded by the scaling.
        report.WriteText(t_strings[16])
        report.Newline()

        report.WriteImage(self.report_images_wx[0])  # insert image buffer in wx rich text
        report.Newline()
        report.Newline()
        report.WriteText(t_strings[17])
        report.Newline()

        report.WriteImage(self.report_images_wx[1])
        report.Newline()
        report.WriteText(t_strings[18])
        report.Newline()
        report.Newline()
        report.BeginFontSize(8)
        report.BeginBold()
        report.WriteText(t_strings[19])
        report.Newline()

        report.WriteImage(self.report_images_wx[2])
        report.Newline()
        report.WriteText(t_strings[20])
        report.Newline()

        report.WriteImage(self.report_images_wx[3])
        report.Newline()
        report.WriteText(t_strings[21])
        report.Newline()

        report.WriteImage(self.report_images_wx[4])
        report.Newline()
        report.EndBold()

    def wordreport(self, docx_file):
        """
        Takes the *t_strings* and *image_files* lists prepared in 'reporter' and writes
        the report to a word document. The document will display the text in the
        default styles in the Word installation used to view the file.  The image
        files retain their native resolution. By default, the name and folder of
        the Excel input file is used simply with .docx added at the end. A folder
        'templates' is assumed to hold default.docx on which to build the report.
        """
        t_strings = self.report_txt
        assert len(t_strings) > 0, "Nothing to report.  Run calculation before attempting to save"
        template = os.path.join(self.cwd, 'templates', 'default.docx')
        document = docx.Document(template)
        document.add_heading(t_strings[0], level=0)
        localtime = time.asctime(time.localtime(time.time()))
        identifier = '(Data source: ' + self.m_statusBar1.GetStatusText(1) + '. Error calculated: ' + localtime + '.)'
        document.add_heading(identifier, level=1)
        document.add_paragraph(t_strings[1] + t_strings[2])
        document.add_paragraph(t_strings[3] + t_strings[4])
        document.add_paragraph(t_strings[5] + t_strings[6])
        document.add_heading(t_strings[16], level=2)
        image_files = self.report_images
        document.add_picture(image_files[0], width=docx.shared.Mm(100))
        document.add_heading(t_strings[17], level=2)
        document.add_picture(image_files[1], width=docx.shared.Mm(100))
        document.add_heading(t_strings[18][:-16] + '.', level=2)
        document.add_heading(t_strings[19] + '.', level=3)
        document.add_picture(image_files[2], width=docx.shared.Mm(100))
        document.add_heading(t_strings[20] + '.', level=3)
        document.add_picture(image_files[3], width=docx.shared.Mm(100))
        document.add_heading(t_strings[21] + '.', level=3)
        document.add_picture(image_files[4], width=docx.shared.Mm(100))
        document.add_page_break()
        document.add_heading('Additional Information', level=1)
        document.add_heading('Uncertainty Contribution Summary', level=2)
        document.add_paragraph(t_strings[7] + ' ' + t_strings[8] + t_strings[9] + ', '
                               + t_strings[10] + t_strings[11] + ', ' + t_strings[12] + t_strings[13] + '.')
        document.add_heading('Text Output From Calculation Process', level=2)
        no_lines = self.m_textCtrl3.GetNumberOfLines()
        for x in range(no_lines):
            output_notes = self.m_textCtrl3.GetLineText(x)
            document.add_paragraph(output_notes)
        document.save(docx_file)

    def OnAutoCalc(self, event):
        self.PushAutocalc()

    def PushAutocalc(self):
        """
        Button push for "Process opened project files" button, proceeds to load all
        data files and then executes 'grand_finale'.
        """
        self.m_statusBar1.SetStatusText('Autocalculation started', 2)
        self.m_statusBar1.SetStatusText('Meter calculation', 0)
        self.PushLoadMeterData()
        self.PushPlotMeter()
        print('Meter data')
        self.PushFit_Meter()
        self.PushLoadMeterInf()
        self.PushMeterSummary()
        self.m_statusBar1.SetStatusText('CT calculation', 0)
        self.PushLoadCTData()
        self.PushPlotCTphase()
        self.PushPlotCTratio()
        print('CT ratio data')
        self.PushFit_CTratio()
        print('CT phase data')
        self.PushFit_CTphase()
        self.PushLoadCTInf()
        self.PushCTSummary()
        self.m_statusBar1.SetStatusText('VT calculation', 0)
        self.PushLoadVTData()
        self.PushPlotVTphase()
        self.PushPlotVTratio()
        print('VT ratio data')
        self.PushFit_VTratio()
        print('VT phase data')
        self.PushFit_VTphase()
        print('')
        self.PushLoadVTInf()
        self.PushVTSummary()
        self.m_statusBar1.SetStatusText('Load processing', 0)
        self.PushPlotLoadProfile()
        self.PushLoadSiteInf()
        self.PushSiteSummary()
        self.m_statusBar1.SetStatusText('Overall calculation', 0)
        self.grand_finale()
    ##########End Manage Calculation & Report###########


if __name__ == '__main__':
    app = wx.App()
    frame = MIECALC(None)
    frame.Show()
    app.MainLoop()
